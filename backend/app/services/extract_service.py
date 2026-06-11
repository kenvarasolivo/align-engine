"""Text extraction for uploaded resume files (PDF, DOCX, TXT).

Pure parsing — no AI involved. The extracted text is returned to the frontend
and placed into the editable resume textarea, keeping the human in the loop.
"""

import io
import os

from docx import Document
from pypdf import PdfReader

MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded file. Raises ValueError on user errors."""
    if not data:
        raise ValueError("The uploaded file is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise ValueError("File is larger than 10 MB.")

    extension = os.path.splitext(filename)[1].lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")

    if extension == ".pdf":
        text = _from_pdf(data)
    elif extension == ".docx":
        text = _from_docx(data)
    else:
        text = _from_txt(data)

    text = text.strip()
    if not text:
        raise ValueError(
            "No readable text found in the file. Scanned/image-only PDFs are not supported — "
            "please paste the text manually."
        )
    return text


def _from_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ValueError(f"Could not read the PDF file: {exc}") from exc
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _from_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ValueError(f"Could not read the DOCX file: {exc}") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _from_txt(data: bytes) -> str:
    try:
        # utf-8-sig strips the BOM that Windows editors prepend to UTF-8 files
        return data.decode("utf-8-sig")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")
