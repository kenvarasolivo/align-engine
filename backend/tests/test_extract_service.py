"""Tests for the non-AI file extraction layer (PDF/DOCX/TXT)."""

import io

import pytest
from docx import Document

from app.services.extract_service import MAX_FILE_BYTES, extract_text


def test_extracts_plain_text():
    text = extract_text("resume.txt", "Hello\nWorld".encode("utf-8"))
    assert text == "Hello\nWorld"


def test_strips_utf8_bom():
    """Windows editors prepend a BOM; it must not leak into the resume text."""
    raw = "Jane Doe".encode("utf-8-sig")  # utf-8-sig prepends the BOM
    assert extract_text("resume.txt", raw) == "Jane Doe"


def test_falls_back_to_latin1_on_bad_utf8():
    # 0xff is invalid UTF-8 but decodable as latin-1.
    text = extract_text("resume.txt", b"caf\xe9")
    assert "caf" in text


def test_rejects_empty_file():
    with pytest.raises(ValueError, match="empty"):
        extract_text("resume.txt", b"")


def test_rejects_oversized_file():
    big = b"x" * (MAX_FILE_BYTES + 1)
    with pytest.raises(ValueError, match="10 MB"):
        extract_text("resume.txt", big)


def test_rejects_unsupported_extension():
    with pytest.raises(ValueError, match="Unsupported"):
        extract_text("resume.rtf", b"some text")


def test_rejects_whitespace_only_after_strip():
    with pytest.raises(ValueError, match="No readable text"):
        extract_text("resume.txt", b"   \n  \t ")


def test_extracts_docx():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Senior Engineer")
    doc.add_paragraph("Built distributed systems")
    doc.save(buf)
    text = extract_text("resume.docx", buf.getvalue())
    assert "Senior Engineer" in text
    assert "distributed systems" in text


def test_corrupt_docx_raises_value_error():
    with pytest.raises(ValueError, match="DOCX"):
        extract_text("resume.docx", b"not really a docx")
